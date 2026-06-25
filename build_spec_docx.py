#!/usr/bin/env python3
"""Build styled DOCX of the Find Your Floor spec, in English and Russian.

Clean header (no Field/Value table), shaded table headers, borders, brand accents.
No long dashes. Run: .venv/bin/python clients/luxfloor/site-assistant/build_spec_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Palette
SLATE = "2C3E50"
BRONZE = "8C6A43"
ROW_ALT = "F5F2EE"
BORDER = "D9D9D9"
WHITE = "FFFFFF"
TEXT = "2B2B2B"
MUTED = "6B6B6B"
BODY_FONT = "Calibri"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), BORDER)
        borders.append(e)
    tblPr.append(borders)


def cell_text(cell, text, bold=False, color=TEXT, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    for i, part in enumerate(text.split("**")):
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = BODY_FONT
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        r.bold = bold or (i % 2 == 1)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    set_borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, SLATE)
        cell_text(c, h, bold=True, color=WHITE, size=10)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            c = cells[j]
            if ri % 2 == 1:
                shade(c, ROW_ALT)
            cell_text(c, str(val), size=10)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BRONZE)
    pbdr.append(bottom)
    pPr.append(pbdr)


def subhead(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BRONZE)


def para(doc, text, italic=False, color=TEXT, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    for i, part in enumerate(text.split("**")):
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = BODY_FONT
        r.font.size = Pt(size)
        r.italic = italic
        r.font.color.rgb = RGBColor.from_string(color)
        r.bold = (i % 2 == 1)


def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    set_borders(t)
    c = t.rows[0].cells[0]
    shade(c, ROW_ALT)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BRONZE)
    r2 = p.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def header_lines(doc, pairs):
    for label, value in pairs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ":  ")
        r.bold = True
        r.font.name = BODY_FONT
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(BRONZE)
        r2 = p.add_run(value)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = BODY_FONT
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(TEXT)


def render(blocks, subtitle, outpath):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('"Find Your Floor"')
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(8)
    rs = sub.add_run(subtitle)
    rs.font.name = BODY_FONT
    rs.font.size = Pt(12)
    rs.font.color.rgb = RGBColor.from_string(BRONZE)

    for b in blocks:
        kind = b[0]
        if kind == "header_lines":
            header_lines(doc, b[1])
        elif kind == "callout":
            callout(doc, b[1], b[2])
        elif kind == "h1":
            h1(doc, b[1])
        elif kind == "subhead":
            subhead(doc, b[1])
        elif kind == "para":
            para(doc, b[1])
        elif kind == "note":
            para(doc, b[1], italic=True, color=MUTED, size=9.5)
        elif kind == "table":
            add_table(doc, b[1], b[2], b[3] if len(b) > 3 else None)
        elif kind == "bullets":
            for x in b[1]:
                bullet(doc, x)
        elif kind == "footer":
            doc.add_paragraph()
            para(doc, b[1], italic=True, color=MUTED, size=9.5)
    doc.save(outpath)
    print("Saved", outpath)


# ---------------- ENGLISH ----------------
EN = [
    ("header_lines", [
        ("Date", "2026-06-08"),
        ("Status", "Draft for review (revisit 2026-06-09 AM to scope the build)"),
        ("Inputs", "Ilya's 4-step qualification flow + scraped site FAQ + CONTEXT.md"),
        ("Gate", "Ilya confirms the cleaned 4-step flow matches how he sells, before any build plan"),
    ]),
    ("callout", "What this is:", "the design for Lux-Floor's site assistant. Not a support bot. A guiding, selling assistant that qualifies a buyer the way Ilya does, explains trade-offs from the FAQ, narrows to products, and captures the lead into the CRM."),
    ("callout", "What this is not:", "a build plan or a stack decision. Those come next, off this spec."),
    ("h1", "1. The core idea"),
    ("para", "Ilya's document revealed the real vision: the assistant should sell, not just answer. It is built as three layers that lock together. The flow asks the question; the FAQ supplies the answer; the output captures the lead."),
    ("table", ["Layer", "Source", "Role"], [
        ["1. Spine", "Ilya's 4-step flow", "Leads the customer through Look, Material, Constraints, Urgency"],
        ["2. Answer bank", "The site FAQ", "Explains every trade-off the flow raises"],
        ["3. Output", "Catalog + Bitrix24", "Recommends products, captures a scored lead"],
    ], [3.2, 4.0, 9.3]),
    ("note", "The two source documents are not competing versions of one thing. They are two different layers of the same assistant."),
    ("h1", "2. Layer 1: the qualification flow (the spine)"),
    ("para", "Ilya's four steps, cleaned up: duplicates removed, options structured. German is customer-facing; English is for our reference."),
    ("subhead", "Step 1: Optik (the look) · soft preferences"),
    ("table", ["Dimension", "Options (DE)", "English"], [
        ["Farbe (Color)", "hell, dunkel, braun, grau", "light, dark, brown, grey"],
        ["Design", "Holzoptik, Steinoptik, Marmoroptik", "wood, stone, marble look"],
        ["Oberfläche (Surface)", "matt, Hochglanz", "matte, high-gloss"],
        ["Kante (Edge)", "mit Fuge, mit Fase, ohne", "with joint, with bevel, none"],
        ["Muster (Pattern)", "Diele, Fischgrät", "plank, herringbone"],
        ["  if herringbone", "Chevron, Englisch", "chevron (45° V), english (90° blocks)"],
    ], [4.0, 6.5, 6.0]),
    ("subhead", "Step 2: Material · type leaning + budget"),
    ("table", ["Dimension", "Options (DE)", "Meaning"], [
        ["Typ-Tendenz (Type)", "dünn, Klick, Echt Holz, wasserfest, robust", "thin / low height, click, real wood, waterproof, durable"],
        ["Preis (Budget)", "range or sensitivity", "which collections to show"],
    ], [4.0, 6.5, 6.0]),
    ("subhead", "Step 3: Probleme nennen (the hard constraints)"),
    ("para", "These rule options in or out. This is where the FAQ does the heavy lifting (see section 3)."),
    ("table", ["Constraint (DE)", "English"], [
        ["Fußbodenheizung vorhanden", "underfloor heating present"],
        ["Alter Belag darf nicht entfernt werden", "old floor stays, install on top"],
        ["Belag darf nicht hoch sein", "low build height required"],
        ["Muss hart sein", "must be hard"],
        ["Muss strapazierfähig sein", "must be durable / high traffic"],
        ["Muss wasserfest sein", "must be waterproof"],
        ["Muss später leicht entfernbar sein", "must be removable later"],
        ["Muss ein Bio-Material sein", "must be eco / natural"],
    ], [8.5, 8.0]),
    ("note", "Deduped from Ilya's notes: \"wasserfest\" appeared in Steps 2 and 3 (kept once, here). \"Fußbodenheizung\" and \"für Fußbodenheizung geeignet\" merged into one."),
    ("subhead", "Step 4: Verfügbarkeit (urgency = the lead signal)"),
    ("table", ["Answer (DE)", "English", "Lead score", "Action"], [
        ["Wird jetzt dringend benötigt", "needed urgently now", "HOT", "real-time alert to Ilya / Alisa"],
        ["Hat Zeit", "has time", "Warm", "normal pipeline"],
        ["Muss eingelagert werden", "must be stored", "Warm + logistics", "note free 30-day storage (FAQ)"],
    ], [4.5, 3.5, 3.0, 5.5]),
    ("h1", "3. Layer 2: constraint to FAQ to recommendation"),
    ("para", "The assistant's reasoning table. Each Step-3 constraint resolves using the FAQ as the answer bank."),
    ("table", ["Constraint", "FAQ fact", "Recommendation"], [
        ["Underfloor heating", "PE-foam underlay OK; cork and polystyrene not for UFH; glue-vinyl ideal", "Vinyl + correct underlay; avoid cork/polystyrene"],
        ["Low height / old floor stays", "Glue-down vinyl has the lowest build height", "Glue-down vinyl, or a thin click floor"],
        ["Removable later", "Glue removal destroys the plank; click is liftable", "Click system, not glue"],
        ["Waterproof", "Vinyl is waterproof; laminate (HDF core) is not", "Vinyl over laminate"],
        ["Durable / high traffic", "Nutzungsklassen 21-43 by room", "Match class to room: 22 living, 23 kitchen/hall, 31-33 commercial"],
        ["Eco / natural", "Cork is 100% recyclable; real wood is natural", "Cork underlay or real-wood parquet"],
        ["Must be hard", "gap", "Confirm with Ilya: scratch resistance or rigidity?"],
    ], [3.8, 6.7, 6.0]),
    ("h1", "4. Layer 3: recommendation + lead capture (the output)"),
    ("subhead", "4.1  Narrowing to products"),
    ("para", "The customer profile maps directly onto filters that already exist on the shop:"),
    ("table", ["Shop filter", "Values"], [
        ["Surface", "Hochglanz, Matt, Strukturiert"],
        ["Format", "Diele, Breitdiele, Fliese, Herringbone, Quadratisch"],
        ["Design", "Holzoptik, Steinoptik, Uni, Marmoroptik"],
    ], [4.0, 12.5]),
    ("para", "So the assistant can hand back a filtered shop link (leanest), or list specific SKUs if connected to the catalog (799 products). Stack decision deferred; both viable."),
    ("subhead", "4.2  Lead capture (the kicker)"),
    ("para", "At the end, the assistant collects name + contact, attaches the filled profile and urgency, and pushes a lead into Bitrix24 (via Alexey's inbound webhook). Urgency becomes the lead score (see Step 4). HOT leads trigger the real-time alert (the Smart Buzz hot-lead alert from the merged vision). Net effect: the assistant is a salesperson and a CRM feeder in one pass. It turns anonymous site traffic into qualified, scored leads."),
    ("subhead", "4.3  Escalation"),
    ("para", "The 6 case-by-case FAQ items (shipping abroad, custom info) and anything outside the flow go to a graceful handoff to a human (phone, WhatsApp, email, showroom)."),
    ("subhead", "4.4  Languages"),
    ("para", "Site is German-only today. v1 ships German first. RU/EN later (logic is language-agnostic; only surface copy changes)."),
    ("h1", "5. Decisions to make tomorrow (2026-06-09 AM)"),
    ("table", ["#", "Decision", "Options / note"], [
        ["1", "\"Hard\" (Step 3) meaning", "scratch resistance vs rigidity (one-line confirm)"],
        ["2", "Catalog connection", "filtered shop links (lean) vs live WooCommerce pull (rich) vs curated SKU map"],
        ["3", "Lead destination + alert channel", "Bitrix24 webhook fields; where HOT alert lands (Telegram / Bitrix / WhatsApp)"],
        ["4", "v1 actions", "recommend only, or also request free sample / book showroom visit"],
        ["5", "v1 channel", "website widget first, WhatsApp second (confirm)"],
        ["6", "Confirm the flow with Ilya", "does the cleaned 4-step flow match how he sells (the gate)"],
    ], [1.0, 5.5, 10.0]),
    ("h1", "6. Open for the build plan (not decided here)"),
    ("bullets", [
        "The stack (Claude-powered widget, kept lean, decided at build).",
        "Hosting, embedding method, the knowledge-base format the widget reads.",
        "How the FAQ + flow are stored and kept current (single source of truth).",
    ]),
    ("footer", "Next: revisit AM 2026-06-09, confirm with Ilya, then create-plan the build off this spec."),
]

# ---------------- RUSSIAN ----------------
RU = [
    ("header_lines", [
        ("Дата", "2026-06-08"),
        ("Статус", "Черновик на ревью (вернуться утром 09.06.2026, расписать сборку)"),
        ("Источники", "4-шаговый сценарий Ильи + FAQ с сайта + CONTEXT.md"),
        ("Контроль", "Илья подтверждает, что очищенный сценарий соответствует тому, как он продаёт, до плана сборки"),
    ]),
    ("callout", "Что это:", "дизайн ассистента для сайта Lux-Floor. Не бот поддержки. Ведущий, продающий ассистент: квалифицирует покупателя так же, как Илья, объясняет компромиссы из FAQ, сужает выбор до товаров и захватывает лид в CRM."),
    ("callout", "Чем это не является:", "это не план сборки и не выбор стека. Они идут следующими, на основе этой спецификации."),
    ("h1", "1. Главная идея"),
    ("para", "Документ Ильи раскрыл настоящее видение: ассистент должен продавать, а не просто отвечать. Он построен из трёх слоёв, которые соединяются вместе. Сценарий задаёт вопрос; FAQ даёт ответ; результат захватывает лид."),
    ("table", ["Слой", "Источник", "Роль"], [
        ["1. Основа", "4-шаговый сценарий Ильи", "Ведёт покупателя через Вид, Материал, Ограничения, Срочность"],
        ["2. База ответов", "FAQ сайта", "Объясняет каждый компромисс, который поднимает сценарий"],
        ["3. Результат", "Каталог + Bitrix24", "Рекомендует товары, захватывает оценённый лид"],
    ], [3.2, 4.5, 8.8]),
    ("note", "Два исходных документа не конкурирующие версии одного. Это два разных слоя одного ассистента."),
    ("h1", "2. Слой 1: сценарий квалификации (основа)"),
    ("para", "Четыре шага Ильи, очищенные: дубликаты убраны, опции структурированы. Немецкий для клиента; русский для нас."),
    ("subhead", "Шаг 1: Optik (внешний вид) · мягкие предпочтения"),
    ("table", ["Параметр", "Опции (DE)", "Русский"], [
        ["Farbe (цвет)", "hell, dunkel, braun, grau", "светлый, тёмный, коричневый, серый"],
        ["Design", "Holzoptik, Steinoptik, Marmoroptik", "под дерево, под камень, под мрамор"],
        ["Oberfläche (поверхность)", "matt, Hochglanz", "матовая, глянец"],
        ["Kante (кромка)", "mit Fuge, mit Fase, ohne", "с фугой, с фаской, без"],
        ["Muster (узор)", "Diele, Fischgrät", "доска, ёлочка"],
        ["  если ёлочка", "Chevron, Englisch", "шеврон (45° V), английская (90° блоки)"],
    ], [4.2, 6.3, 6.0]),
    ("subhead", "Шаг 2: Material · склонность по типу + бюджет"),
    ("table", ["Параметр", "Опции (DE)", "Значение"], [
        ["Typ-Tendenz (тип)", "dünn, Klick, Echt Holz, wasserfest, robust", "тонкий / малая высота, клик, настоящее дерево, водостойкий, прочный"],
        ["Preis (бюджет)", "диапазон или чувствительность", "какие коллекции показывать"],
    ], [4.2, 6.3, 6.0]),
    ("subhead", "Шаг 3: Probleme nennen (жёсткие ограничения)"),
    ("para", "Они включают или исключают опции. Здесь основную работу делает FAQ (см. раздел 3)."),
    ("table", ["Ограничение (DE)", "Русский"], [
        ["Fußbodenheizung vorhanden", "тёплый пол присутствует"],
        ["Alter Belag darf nicht entfernt werden", "старое покрытие нельзя снимать, монтаж поверх"],
        ["Belag darf nicht hoch sein", "нужна малая высота покрытия"],
        ["Muss hart sein", "должно быть твёрдым"],
        ["Muss strapazierfähig sein", "должно быть износостойким / высокий трафик"],
        ["Muss wasserfest sein", "должно быть водостойким"],
        ["Muss später leicht entfernbar sein", "должно легко сниматься позже"],
        ["Muss ein Bio-Material sein", "должно быть эко / натуральным"],
    ], [8.5, 8.0]),
    ("note", "Убраны дубликаты из заметок Ильи: «wasserfest» был в шагах 2 и 3 (оставлено один раз, здесь). «Fußbodenheizung» и «für Fußbodenheizung geeignet» объединены."),
    ("subhead", "Шаг 4: Verfügbarkeit (срочность = сигнал лида)"),
    ("table", ["Ответ (DE)", "Русский", "Оценка лида", "Действие"], [
        ["Wird jetzt dringend benötigt", "нужно срочно сейчас", "ГОРЯЧИЙ", "мгновенный алерт Илье / Алисе"],
        ["Hat Zeit", "есть время", "Тёплый", "обычная воронка"],
        ["Muss eingelagert werden", "нужно хранение", "Тёплый + логистика", "отметить бесплатное хранение 30 дней (FAQ)"],
    ], [4.5, 3.5, 3.0, 5.5]),
    ("h1", "3. Слой 2: ограничение, FAQ, рекомендация"),
    ("para", "Таблица рассуждений ассистента. Каждое ограничение из шага 3 решается через FAQ как базу ответов."),
    ("table", ["Ограничение", "Факт из FAQ", "Рекомендация"], [
        ["Тёплый пол", "PE-подложка ок; пробка и пенополистирол не для тёплого пола; клеевой винил идеален", "Винил + правильная подложка; избегать пробки/пенополистирола"],
        ["Малая высота / старый пол остаётся", "У клеевого винила наименьшая высота", "Клеевой винил или тонкий клик-пол"],
        ["Снять позже", "Клей разрушает планку при снятии; клик снимается", "Клик-система, не клей"],
        ["Водостойкость", "Винил водостойкий; ламинат (ядро HDF) нет", "Винил вместо ламината"],
        ["Износостойкость / трафик", "Классы износа 21-43 по помещению", "Класс под помещение: 22 гостиная, 23 кухня/прихожая, 31-33 коммерция"],
        ["Эко / натуральное", "Пробка на 100% перерабатываема; дерево натурально", "Пробковая подложка или паркет из настоящего дерева"],
        ["Должно быть твёрдым", "пробел", "Уточнить у Ильи: стойкость к царапинам или жёсткость?"],
    ], [3.8, 6.7, 6.0]),
    ("h1", "4. Слой 3: рекомендация + захват лида (результат)"),
    ("subhead", "4.1  Сужение до товаров"),
    ("para", "Профиль клиента напрямую ложится на фильтры, которые уже есть в магазине:"),
    ("table", ["Фильтр магазина", "Значения"], [
        ["Surface (поверхность)", "Hochglanz, Matt, Strukturiert"],
        ["Format (формат)", "Diele, Breitdiele, Fliese, Herringbone, Quadratisch"],
        ["Design (дизайн)", "Holzoptik, Steinoptik, Uni, Marmoroptik"],
    ], [4.5, 12.0]),
    ("para", "Ассистент может вернуть ссылку на отфильтрованный магазин (самое лёгкое) или перечислить конкретные SKU, если подключить каталог (799 товаров). Выбор стека отложен; оба варианта рабочие."),
    ("subhead", "4.2  Захват лида (ключевой момент)"),
    ("para", "В конце ассистент собирает имя + контакт, прикрепляет заполненный профиль и срочность и отправляет лид в Bitrix24 (через входящий вебхук Алексея). Срочность становится оценкой лида (см. шаг 4). Горячие лиды запускают мгновенный алерт (это алерт горячих лидов Smart Buzz из общего видения). Итог: ассистент одновременно и продавец, и поставщик лидов в CRM. Он превращает анонимный трафик сайта в квалифицированные, оценённые лиды."),
    ("subhead", "4.3  Эскалация"),
    ("para", "6 пунктов FAQ «по ситуации» (доставка за границу, особые вопросы) и всё вне сценария идут на мягкую передачу человеку (телефон, WhatsApp, email, шоурум)."),
    ("subhead", "4.4  Языки"),
    ("para", "Сайт сегодня только на немецком. v1 выходит сначала на немецком. RU/EN позже (логика не зависит от языка; меняется только текст интерфейса)."),
    ("h1", "5. Решения на завтра (утро 09.06.2026)"),
    ("table", ["№", "Решение", "Опции / заметка"], [
        ["1", "Значение «твёрдый» (шаг 3)", "стойкость к царапинам или жёсткость (подтвердить одной строкой)"],
        ["2", "Подключение каталога", "ссылки на фильтр магазина (лёгкое) vs живой запрос WooCommerce (богаче) vs ручная карта SKU"],
        ["3", "Назначение лида + канал алерта", "поля вебхука Bitrix24; куда идёт ГОРЯЧИЙ алерт (Telegram / Bitrix / WhatsApp)"],
        ["4", "Действия v1", "только рекомендация, или ещё запрос образца / запись в шоурум"],
        ["5", "Канал v1", "сначала виджет на сайте, WhatsApp вторым (подтвердить)"],
        ["6", "Подтвердить сценарий с Ильёй", "соответствует ли очищенный сценарий тому, как он продаёт (контроль)"],
    ], [1.0, 5.5, 10.0]),
    ("h1", "6. Открыто для плана сборки (здесь не решается)"),
    ("bullets", [
        "Стек (виджет на Claude, держим лёгким, решаем при сборке).",
        "Хостинг, способ встраивания, формат базы знаний, которую читает виджет.",
        "Как FAQ + сценарий хранятся и поддерживаются в актуальности (единый источник правды).",
    ]),
    ("footer", "Далее: вернуться утром 09.06.2026, подтвердить с Ильёй, затем create-plan сборки на основе этой спецификации."),
]


if __name__ == "__main__":
    base = "clients/luxfloor/site-assistant/"
    render(EN, "Site Assistant Spec  ·  v1 Design", base + "find-your-floor-spec.docx")
    render(RU, "Спецификация ассистента сайта  ·  Дизайн v1", base + "find-your-floor-spec-ru.docx")
