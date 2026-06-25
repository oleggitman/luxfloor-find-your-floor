#!/usr/bin/env python3
"""
Build DOCX: deployment task for Alexey.
Run: .venv/bin/python clients/luxfloor/site-assistant/build_deploy_task_alexey.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SLATE  = "2C3E50"
BRONZE = "8C6A43"
ROW_ALT = "F5F2EE"
BORDER = "D9D9D9"
WHITE  = "FFFFFF"
TEXT   = "2B2B2B"
MUTED  = "6B6B6B"
GREEN  = "27AE60"
FONT   = "Calibri"


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), BORDER)
        borders.append(e)
    tblPr.append(borders)


def cell_text(cell, text, bold=False, color=TEXT, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, SLATE)
        cell_text(c, h, bold=True, color=WHITE, size=10)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for j, val in enumerate(row):
            c = cells[j]
            if ri % 2 == 1:
                shade(c, ROW_ALT)
            cell_text(c, str(val), size=10)
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(13)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(SLATE)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BRONZE)
    pbdr.append(bottom)
    pPr.append(pbdr)


def step_head(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"Шаг {number}.  ")
    r1.font.name = FONT
    r1.font.size = Pt(11)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(BRONZE)
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(11)
    r2.bold = True
    r2.font.color.rgb = RGBColor.from_string(SLATE)


def para(doc, text, italic=False, color=TEXT, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + "  ")
        r1.font.name = FONT
        r1.font.size = Pt(10.5)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(TEXT)
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor.from_string(TEXT)
    else:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor.from_string(TEXT)


def callout(doc, label, text, color=ROW_ALT):
    t = doc.add_table(rows=1, cols=1)
    set_borders(t)
    c = t.rows[0].cells[0]
    shade(c, color)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BRONZE)
    r2 = p.add_run(text)
    r2.font.name = FONT
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def code_block(doc, text):
    t = doc.add_table(rows=1, cols=1)
    set_borders(t)
    c = t.rows[0].cells[0]
    shade(c, "1E1E1E")
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("D4D4D4")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    for s in doc.sections:
        s.top_margin    = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin   = Cm(2.2)
        s.right_margin  = Cm(2.2)

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Find Your Floor")
    r.bold = True
    r.font.name = FONT
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    rs = sub.add_run("Запуск ассистента на сайте: задание для Алексея")
    rs.font.name = FONT
    rs.font.size = Pt(12)
    rs.font.color.rgb = RGBColor.from_string(BRONZE)

    # Meta
    meta_t = doc.add_table(rows=3, cols=2)
    set_borders(meta_t)
    meta_data = [
        ("Дата", "2026-06-24"),
        ("Время на выполнение", "около 20-30 минут"),
        ("Вопросы", "писать Олегу"),
    ]
    for i, (label, value) in enumerate(meta_data):
        c0 = meta_t.rows[i].cells[0]
        c1 = meta_t.rows[i].cells[1]
        shade(c0, ROW_ALT)
        cell_text(c0, label, bold=True, color=BRONZE, size=10)
        cell_text(c1, value, size=10)
    for row in meta_t.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Intro
    h1(doc, "Что это")
    para(doc,
         "На сайте lux-floor.de появится кнопка-ассистент. Посетитель нажимает, "
         "вводит вопрос или запрос, ассистент задаёт уточняющие вопросы, подбирает "
         "пол из каталога и создаёт лид в Bitrix24. Всё работает автоматически. "
         "Тебе нужно один раз развернуть сервер и вставить одну строку в сайт.")

    callout(doc, "Что Олег уже сделал:",
            "написал весь код, проверил на живом каталоге, "
            "подготовил конфигурацию сервера. Тебе ничего настраивать не нужно, "
            "только нажать несколько кнопок и вставить ключи.")

    # Step 1
    h1(doc, "Шаги")
    step_head(doc, 1, "Создать аккаунт на Render (2 минуты)")
    bullet(doc, "render.com")
    bullet(doc, 'Нажать "Get Started for free"')
    bullet(doc, "Зарегистрироваться через GitHub (или email)")
    bullet(doc, "Подтвердить email если просит")
    para(doc, "Render — это сервер-хостинг, на котором будет жить ассистент. "
              "Бесплатная регистрация.", color=MUTED, size=9.5, italic=True)

    # Step 2
    step_head(doc, 2, "Создать новый веб-сервис (3 минуты)")
    bullet(doc, 'В дашборде нажать "New" > "Web Service"')
    bullet(doc, '"Connect a repository" > вставить ссылку на репозиторий (Олег пришлёт)')
    bullet(doc, "Render сам найдёт все настройки из файла в репо. Ничего не менять.")
    bullet(doc, '"Create Web Service"')
    para(doc, "Render автоматически читает файл render.yaml из репозитория и "
              "знает как собрать и запустить сервис.", color=MUTED, size=9.5, italic=True)

    # Step 3
    step_head(doc, 3, "Ввести секретные ключи (5 минут)")
    para(doc, "В настройках сервиса: Environment > Environment Variables. "
              "Добавить 4 ключа. Значения Олег пришлёт отдельным сообщением (не в этом файле).")
    add_table(doc,
        ["Название ключа", "Что это"],
        [
            ["ANTHROPIC_API_KEY",   "Ключ Claude AI (мозг ассистента)"],
            ["WOO_CONSUMER_KEY",    "Ключ WooCommerce (доступ к каталогу)"],
            ["WOO_CONSUMER_SECRET", "Секрет WooCommerce"],
            ["BITRIX_WEBHOOK_URL",  "Вебхук Bitrix24 (ты уже его создавал)"],
        ],
        [6.5, 10.0],
    )
    callout(doc, "Важно:",
            "ключи вводить точно как написано, без пробелов до и после знака =. "
            "Не пересылать их в открытых чатах.")

    # Step 4
    step_head(doc, 4, "Перейти на Starter план - $7 в месяц (2 минуты)")
    bullet(doc, "Billing > Upgrade to Starter")
    bullet(doc, "Ввести карту (Ильи или рабочую)")
    para(doc, "На бесплатном плане сервер засыпает через 15 минут без посетителей. "
              "Первый посетитель после сна ждёт 30 секунд. "
              "Starter ($7/мес) держит сервер живым постоянно.", color=MUTED, size=9.5, italic=True)

    # Step 5
    step_head(doc, 5, "Скопировать URL и отправить Олегу (1 минута)")
    bullet(doc, "После деплоя Render покажет URL сервиса")
    bullet(doc, "Вид: find-your-floor.onrender.com (или похожий)")
    bullet(doc, "Скопировать и отправить Олегу")
    para(doc, "Олег проверит что ассистент отвечает, и пришлёт готовую строку для вставки в сайт.",
         color=MUTED, size=9.5, italic=True)

    # Step 6
    step_head(doc, 6, "Вставить код в WordPress (5 минут) - после подтверждения от Олега")
    para(doc, "После того как Олег подтвердит, что ассистент работает:")

    bullet(doc, "Вариант A (проще): установить плагин",
           bold_prefix=None)
    para(doc, 'WordPress > Plugins > Add New > искать "Insert Headers and Footers" '
              '(автор WPBeginner) > Install > Activate.',
         size=10)

    bullet(doc, "Вариант B (без плагина):")
    para(doc, "Appearance > Theme Editor > footer.php",
         size=10)

    para(doc, "В обоих случаях вставить одну строку кода перед тегом </body>. "
              "Строку пришлёт Олег после шага 5.")

    code_block(doc,
               '<script src="https://find-your-floor.onrender.com/widget.js"\n'
               '        data-backend="https://find-your-floor.onrender.com"></script>\n'
               '(точный URL Олег подтвердит)')

    para(doc, "После вставки: открыть lux-floor.de в браузере, убедиться что "
              "в правом нижнем углу появилась кнопка ассистента.")

    # Summary
    h1(doc, "Итог")
    add_table(doc,
        ["Шаг", "Действие", "Время"],
        [
            ["1", "Регистрация на Render",               "2 мин"],
            ["2", "Создать Web Service, подключить репо", "3 мин"],
            ["3", "Ввести 4 ключа",                      "5 мин"],
            ["4", "Перейти на Starter ($7/мес)",          "2 мин"],
            ["5", "Отправить URL Олегу",                  "1 мин"],
            ["6", "Вставить строку в WordPress",          "5 мин"],
        ],
        [1.0, 9.5, 2.5],
    )

    callout(doc, "Вопросы:", "писать Олегу в любое время.")

    out = "clients/luxfloor/site-assistant/alexey-deploy-task.docx"
    doc.save(out)
    print("Saved:", out)


if __name__ == "__main__":
    build()
